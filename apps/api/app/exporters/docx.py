"""
ResumeForge AI - Word DOCX Document Generator
Generates clean, ATS-compliant Word DOCX resumes using python-docx.
"""

import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.schemas.resume import StructuredResumeContent


def generate_docx_resume(
    resume: StructuredResumeContent,
    template_name: str = "classic",
) -> bytes:
    """Generate ATS-friendly DOCX bytes from structured resume content."""
    doc = Document()

    # Set 0.5 inch margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Accent color based on template
    tpl = template_name.lower()
    if tpl == "professional":
        accent_rgb = RGBColor(30, 58, 138)
    elif tpl == "modern":
        accent_rgb = RGBColor(15, 118, 110)
    else:
        accent_rgb = RGBColor(17, 24, 39)

    # Helper: Add Section Header
    def add_section_header(title: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = accent_rgb

    # 1. Header (Name & Contact)
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_after = Pt(2)
    r_name = p_name.add_run(resume.personal.name or "Your Name")
    r_name.bold = True
    r_name.font.size = Pt(18)
    r_name.font.color.rgb = accent_rgb

    if resume.personal.title:
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.space_after = Pt(3)
        r_title = p_title.add_run(resume.personal.title)
        r_title.bold = True
        r_title.font.size = Pt(11)
        r_title.font.color.rgb = RGBColor(75, 85, 99)

    # Contact line
    contact_parts = []
    if resume.personal.email:
        contact_parts.append(resume.personal.email)
    if resume.personal.phone:
        contact_parts.append(resume.personal.phone)
    if resume.personal.location:
        contact_parts.append(resume.personal.location)
    if resume.personal.linkedin:
        contact_parts.append("LinkedIn")
    if resume.personal.github:
        contact_parts.append("GitHub")

    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_after = Pt(8)
    r_contact = p_contact.add_run(" | ".join(contact_parts))
    r_contact.font.size = Pt(9.5)
    r_contact.font.color.rgb = RGBColor(107, 114, 128)

    # 2. Professional Summary
    if resume.summary:
        add_section_header("Professional Summary")
        p_sum = doc.add_paragraph()
        p_sum.paragraph_format.space_after = Pt(6)
        r_sum = p_sum.add_run(resume.summary)
        r_sum.font.size = Pt(10)

    # 3. Technical Skills
    if resume.skills:
        add_section_header("Technical Skills")
        for cat in resume.skills:
            if cat.items:
                p_sk = doc.add_paragraph()
                p_sk.paragraph_format.space_after = Pt(2)
                r_cat = p_sk.add_run(f"{cat.category}: ")
                r_cat.bold = True
                r_cat.font.size = Pt(9.5)
                r_items = p_sk.add_run(", ".join(cat.items))
                r_items.font.size = Pt(9.5)

    # 4. Work Experience
    if resume.experience:
        add_section_header("Work Experience")
        for exp in resume.experience:
            p_exp = doc.add_paragraph()
            p_exp.paragraph_format.space_before = Pt(4)
            p_exp.paragraph_format.space_after = Pt(1)
            
            r_pos = p_exp.add_run(exp.position)
            r_pos.bold = True
            r_pos.font.size = Pt(10.5)

            date_str = f"{exp.start_date} – {exp.end_date or ('Present' if exp.is_current else '')}"
            r_date = p_exp.add_run(f" | {date_str}")
            r_date.font.size = Pt(9.5)
            r_date.font.color.rgb = RGBColor(107, 114, 128)

            loc_str = f" | {exp.location}" if exp.location else ""
            p_comp = doc.add_paragraph()
            p_comp.paragraph_format.space_after = Pt(2)
            r_comp = p_comp.add_run(f"{exp.company}{loc_str}")
            r_comp.italic = True
            r_comp.font.size = Pt(9.5)

            for h in exp.highlights:
                p_hl = doc.add_paragraph(style="List Bullet")
                p_hl.paragraph_format.space_after = Pt(1)
                r_hl = p_hl.add_run(h)
                r_hl.font.size = Pt(9.5)

    # 5. Key Projects
    if resume.projects:
        add_section_header("Key Projects")
        for proj in resume.projects:
            p_proj = doc.add_paragraph()
            p_proj.paragraph_format.space_before = Pt(4)
            p_proj.paragraph_format.space_after = Pt(1)
            
            r_title = p_proj.add_run(proj.title)
            r_title.bold = True
            r_title.font.size = Pt(10)
            
            if proj.technologies:
                r_tech = p_proj.add_run(f" ({', '.join(proj.technologies)})")
                r_tech.italic = True
                r_tech.font.size = Pt(9)

            if proj.description:
                p_desc = doc.add_paragraph()
                p_desc.paragraph_format.space_after = Pt(1)
                r_desc = p_desc.add_run(proj.description)
                r_desc.font.size = Pt(9.5)

            for h in proj.highlights:
                p_hl = doc.add_paragraph(style="List Bullet")
                p_hl.paragraph_format.space_after = Pt(1)
                r_hl = p_hl.add_run(h)
                r_hl.font.size = Pt(9.5)

    # 6. Education
    if resume.education:
        add_section_header("Education")
        for edu in resume.education:
            p_edu = doc.add_paragraph()
            p_edu.paragraph_format.space_before = Pt(3)
            p_edu.paragraph_format.space_after = Pt(1)

            r_deg = p_edu.add_run(edu.degree)
            r_deg.bold = True
            r_deg.font.size = Pt(10)

            date_str = f"{edu.start_date or ''} – {edu.end_date or ''}".strip(" – ")
            if date_str:
                r_dt = p_edu.add_run(f" | {date_str}")
                r_dt.font.size = Pt(9.5)
                r_dt.font.color.rgb = RGBColor(107, 114, 128)

            p_inst = doc.add_paragraph()
            p_inst.paragraph_format.space_after = Pt(2)
            honors_str = f" | {', '.join(edu.honors)}" if edu.honors else ""
            r_inst = p_inst.add_run(f"{edu.institution}{honors_str}")
            r_inst.font.size = Pt(9.5)

    # 7. Certifications
    if resume.certifications:
        add_section_header("Certifications")
        for c in resume.certifications:
            p_cert = doc.add_paragraph(style="List Bullet")
            p_cert.paragraph_format.space_after = Pt(1)
            issuer = f" ({c.issuer})" if c.issuer else ""
            r_c = p_cert.add_run(f"{c.name}{issuer}")
            r_c.font.size = Pt(9.5)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
