"""
ResumeForge AI - DOCX Document Parser
Extracts plain text, structured paragraphs, lists, and tables from Word DOCX files.
"""

import io
from typing import List
from docx import Document
from app.core.exceptions import DocumentParsingException
from app.parsers.text_cleaner import clean_text


class DOCXParser:
    """Parses Word DOCX documents."""

    @staticmethod
    def extract_text(docx_bytes: bytes) -> str:
        """Extract full plain text from DOCX bytes."""
        if not docx_bytes:
            raise DocumentParsingException("Empty DOCX file provided.")

        try:
            doc = Document(io.BytesIO(docx_bytes))
            content_chunks: List[str] = []

            # Extract paragraphs and headings
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                # If it's a list item style, ensure a bullet prefix
                style_name = (para.style.name or "").lower()
                if "list" in style_name or "bullet" in style_name:
                    if not text.startswith("- ") and not text.startswith("* "):
                        text = f"- {text}"
                content_chunks.append(text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    # Deduplicate adjacent cells if merged
                    unique_cells = []
                    for c in row_cells:
                        if not unique_cells or unique_cells[-1] != c:
                            unique_cells.append(c)
                    if unique_cells:
                        content_chunks.append(" | ".join(unique_cells))

            raw_text = "\n".join(content_chunks)
        except Exception as e:
            raise DocumentParsingException(
                message=f"Failed to parse DOCX document: {str(e)}",
                details={"parser_error": str(e)},
            )

        cleaned = clean_text(raw_text)
        if not cleaned:
            raise DocumentParsingException(
                message="Could not extract readable text from DOCX document.",
                details={"reason": "empty_extracted_text"},
            )

        return cleaned
