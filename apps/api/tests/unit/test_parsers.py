"""
ResumeForge AI - Parser Unit Tests
"""

import io
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from app.parsers.docx_parser import DOCXParser
from app.parsers.pdf_parser import PDFParser
from app.parsers.section_extractor import parse_resume_sections
from app.parsers.text_cleaner import clean_text


def test_text_cleaner():
    dirty = "Hello\u00A0World!\r\n• Point 1\r\n\r\n\r\n▪ Point 2   with   spaces."
    cleaned = clean_text(dirty)
    assert "Hello World!" in cleaned
    assert "- Point 1" in cleaned
    assert "- Point 2 with spaces." in cleaned
    assert "\r" not in cleaned


def test_section_extractor(sample_resume_text: str):
    parsed = parse_resume_sections(sample_resume_text)
    assert parsed.personal.name == "Alex Mercer"
    assert parsed.personal.email == "alex.mercer@example.com"
    assert parsed.personal.phone == "(555) 123-4567"
    assert "San Francisco" in (parsed.personal.location or "")
    assert len(parsed.experience) >= 2
    assert len(parsed.skills) >= 1
    assert len(parsed.education) >= 1
    assert "Berkeley" in parsed.education[0].institution or "Computer Science" in parsed.education[0].degree


def test_docx_parser():
    # Build in-memory docx
    doc = Document()
    doc.add_heading("John Doe", 0)
    doc.add_paragraph("john@example.com | 123-456-7890")
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("Built scalable microservices")
    
    bio = io.BytesIO()
    doc.save(bio)
    docx_bytes = bio.getvalue()

    extracted = DOCXParser.extract_text(docx_bytes)
    assert "John Doe" in extracted
    assert "john@example.com" in extracted
    assert "Built scalable microservices" in extracted


def test_pdf_parser():
    # Build in-memory PDF
    bio = io.BytesIO()
    c = canvas.Canvas(bio, pagesize=letter)
    c.drawString(100, 750, "Jane Doe")
    c.drawString(100, 730, "jane@example.com | Software Engineer")
    c.drawString(100, 710, "- Experienced with Python and Docker")
    c.save()
    pdf_bytes = bio.getvalue()

    extracted = PDFParser.extract_text(pdf_bytes)
    assert "Jane Doe" in extracted
    assert "jane@example.com" in extracted
    assert "Python" in extracted
